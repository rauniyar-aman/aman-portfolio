import base64
import binascii
import io
import os
import re
import sys

from django.template.loader import render_to_string
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DEFAULT_THEME_COLOR = "#1F4E37"
DEFAULT_PHONE_COUNTRY_CODE = "+977"
TEMPLATE_CHOICES = ("classic", "modern", "minimalist", "academic")


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.lstrip("#").upper())


def _clean_link(raw: str) -> dict:
    """Turn a user-entered URL (with or without a scheme) into a normalized
    href plus the bare "domain/path" text the CV should display, matching
    how contact links conventionally look on a printed CV.
    """
    raw = raw.strip()
    href = raw if re.match(r"^https?://", raw, re.IGNORECASE) else f"https://{raw}"
    text = re.sub(r"^https?://(www\.)?", "", raw, flags=re.IGNORECASE).rstrip("/")
    return {"href": href, "text": text}


def build_cv_view_model(cv) -> dict:
    """Shape a CV's raw JSON content into the exact fields/sections the PDF
    and DOCX exporters render, including which optional sections should be
    omitted. Shared by both exporters (and all four template styles) so
    their omission rules can't drift.
    """
    content = cv.content or {}
    theme_color = content.get("theme_color") or DEFAULT_THEME_COLOR

    template = content.get("template") or "classic"
    if template not in TEMPLATE_CHOICES:
        template = "classic"

    photo = content.get("photo") or ""

    full_name = content.get("full_name", "")
    email = content.get("email", "")
    phone = content.get("phone", "")
    phone_country_code = content.get("phone_country_code") or (
        DEFAULT_PHONE_COUNTRY_CODE if phone else ""
    )
    phone_display = f"{phone_country_code} {phone}".strip() if phone else ""
    summary = content.get("summary", "")
    dob = content.get("dob", "")
    nationality = content.get("nationality", "")
    marital_status = content.get("marital_status", "")
    passport = content.get("passport") or {}
    education = content.get("education") or []
    experience = content.get("experience") or []
    references = content.get("references") or []

    # Free text now — the user writes it however they like (one per line,
    # comma-separated, a short paragraph). Older CVs may still have it saved
    # as a list from before that change, so normalize either shape.
    raw_skills = content.get("skills") or ""
    skills = "\n".join(s for s in raw_skills if s) if isinstance(raw_skills, list) else raw_skills

    # Address used to be a plain string; it's now {detail, postcode, country}
    # so the country can be a dropdown and the postcode (never guessed by
    # the passport scanner — entered manually) its own field. Older CVs may
    # still have the plain-string shape, so normalize both into one display
    # string; the postcode segment is simply omitted when blank, never
    # rendered as an empty/placeholder slot.
    raw_address = content.get("address") or ""
    if isinstance(raw_address, dict):
        address = ", ".join(
            part
            for part in (
                raw_address.get("detail", ""),
                raw_address.get("postcode", ""),
                raw_address.get("country", ""),
            )
            if part
        )
    else:
        address = raw_address

    # {language, proficiency} objects, rendered one per bullet as
    # "{language} - {proficiency}". Older CVs saved before proficiency
    # levels existed may still have plain strings here — those render as
    # just the bare name rather than being dropped.
    raw_languages = content.get("languages") or []
    language_lines = []
    for lang in raw_languages:
        if isinstance(lang, dict):
            name = str(lang.get("language", "")).strip()
            proficiency = str(lang.get("proficiency", "")).strip()
        elif isinstance(lang, str):
            name = lang.strip()
            proficiency = ""
        else:
            continue
        if not name:
            continue
        language_lines.append(f"{name} - {proficiency}" if proficiency else name)

    # Only used by the academic template; omitted (no heading) if empty.
    raw_publications = content.get("publications") or []
    publication_rows = [
        {
            "title": pub.get("title", ""),
            "venue": pub.get("venue", ""),
            "year": pub.get("year", ""),
        }
        for pub in raw_publications
        if pub.get("title")
    ]

    # Name, address, phone/email, and links all render in the header block
    # (see the reference layout), not as body sections — so Personal Details
    # only carries DOB/Nationality/Marital Status now.
    personal_rows = [
        row
        for row in (
            ("Date of Birth", dob),
            ("Nationality", nationality),
            ("Marital Status", marital_status),
        )
        if row[1]
    ]

    link_items = [_clean_link(link) for link in (content.get("links") or []) if link and link.strip()]

    passport_rows = [
        row
        for row in (
            ("Passport Number", passport.get("number", "")),
            ("Issued By", passport.get("issued_by", "")),
            ("Issued Date", passport.get("issued_date", "")),
            ("Expiry Date", passport.get("expiry_date", "")),
        )
        if row[1]
    ]

    # Four purpose-agnostic optional sections (most relevant for foreign
    # employment CVs, but available regardless of purpose) — each omitted
    # entirely unless it actually has data.
    physical_details = content.get("physical_details") or {}
    physical_details_rows = [
        row
        for row in (
            ("Height", physical_details.get("height", "")),
            ("Weight", physical_details.get("weight", "")),
            ("Blood Group", physical_details.get("blood_group", "")),
        )
        if row[1]
    ]

    emergency_contact = content.get("emergency_contact") or {}
    emergency_contact_rows = [
        row
        for row in (
            ("Name", emergency_contact.get("name", "")),
            ("Relationship", emergency_contact.get("relationship", "")),
            ("Phone", emergency_contact.get("phone", "")),
        )
        if row[1]
    ]

    preferred_position = content.get("preferred_position", "")
    medical_fitness = content.get("medical_fitness", "")

    include_declaration = bool(content.get("include_declaration"))
    declaration_text = content.get("declaration_text") or ""
    declaration_date = content.get("declaration_date") or ""

    education_rows = [
        {
            "degree": item.get("degree", ""),
            "institute": item.get("institute", ""),
            "address": item.get("address", ""),
            "percentage_grade": item.get("percentage_grade", ""),
            "duration": " - ".join(
                d for d in (item.get("start_date", ""), item.get("end_date", "")) if d
            ),
        }
        for item in education
    ]

    experience_rows = [
        {
            "company": item.get("company", ""),
            "position": item.get("position", ""),
            "tenure": " – ".join(
                d
                for d in (item.get("start_date", ""), item.get("end_date_or_present", ""))
                if d
            ),
            "responsibilities": [r for r in (item.get("responsibilities") or []) if r],
        }
        for item in experience
    ]

    # Omitted entirely (no heading, no rows) unless the user added at least
    # one — references are the one section that's genuinely optional on a CV.
    reference_rows = [
        {
            "name": item.get("name", ""),
            "position": item.get("position", ""),
            "company": item.get("company", ""),
            "phone": item.get("phone", ""),
            "email": item.get("email", ""),
        }
        for item in references
        if item.get("name")
    ]

    return {
        "template": template,
        "theme_color": theme_color,
        "photo": photo,
        "full_name": full_name,
        "address": address,
        "phone": phone_display,
        "email": email,
        "link_items": link_items,
        "summary": summary,
        "personal_rows": personal_rows,
        "passport_rows": passport_rows,
        "physical_details_rows": physical_details_rows,
        "emergency_contact_rows": emergency_contact_rows,
        "preferred_position": preferred_position,
        "medical_fitness": medical_fitness,
        "education_rows": education_rows,
        "experience_rows": experience_rows,
        "skills": skills,
        "reference_rows": reference_rows,
        "language_lines": language_lines,
        "publication_rows": publication_rows,
        "include_declaration": include_declaration,
        "declaration_text": declaration_text,
        "declaration_date": declaration_date,
    }


_GTK_DLL_DIR_ADDED = False


def _ensure_windows_gtk_dll_dir():
    """WeasyPrint locates its native GTK/Pango DLLs via the process's DLL
    search path. On Windows, installers register the runtime's bin/ folder
    in the *system* PATH, but already-running processes (this one included,
    if it was started before the install, or a shell that hasn't been
    restarted since) never see that update — Windows doesn't broadcast PATH
    changes to running processes. Adding the known install directory
    explicitly here makes PDF export work regardless of when/how the
    process was started, instead of depending on a shell restart or reboot.
    """
    global _GTK_DLL_DIR_ADDED
    if _GTK_DLL_DIR_ADDED or sys.platform != "win32":
        return

    for candidate in (
        r"C:\Program Files\GTK3-Runtime Win64\bin",
        r"C:\Program Files (x86)\GTK3-Runtime Win64\bin",
    ):
        if os.path.isdir(candidate):
            if hasattr(os, "add_dll_directory"):
                os.add_dll_directory(candidate)
            os.environ["PATH"] = candidate + os.pathsep + os.environ.get("PATH", "")
            _GTK_DLL_DIR_ADDED = True
            return


PDF_TEMPLATES = {
    "classic": "cvs/cv_pdf.html",
    "modern": "cvs/cv_pdf_modern.html",
    "minimalist": "cvs/cv_pdf_minimalist.html",
    "academic": "cvs/cv_pdf_academic.html",
}


def render_cv_pdf(cv) -> bytes:
    # Imported lazily: WeasyPrint requires native GTK/Pango libraries that may
    # not be installed on every host, and we don't want that to break imports
    # for the rest of the API (see WeasyPrint's Windows install docs).
    _ensure_windows_gtk_dll_dir()
    from weasyprint import HTML

    vm = build_cv_view_model(cv)
    template_name = PDF_TEMPLATES.get(vm["template"], PDF_TEMPLATES["classic"])
    html_string = render_to_string(template_name, vm)
    return HTML(string=html_string).write_pdf()


def _add_bottom_border(paragraph, color_hex: str, size: int = 8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color_hex.lstrip("#").upper())
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_hyperlink(paragraph, url: str, text: str, color_hex: str):
    """python-docx has no built-in hyperlink support; this is the standard
    OOXML recipe: register an external relationship, then build a
    <w:hyperlink> run pointing at it."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex.lstrip("#").upper())
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_cell_background(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex.lstrip("#").upper())
    tcPr.append(shd)


def _decode_photo(photo_data_uri: str):
    if not photo_data_uri or "," not in photo_data_uri:
        return None
    _, encoded = photo_data_uri.split(",", 1)
    try:
        return io.BytesIO(base64.b64decode(encoded))
    except (ValueError, TypeError, binascii.Error):
        return None


def _make_picture_circular(inline_shape):
    """python-docx has no crop/clip API; Word (and most viewers) render an
    inline picture circularly if its geometry is overridden to "ellipse" on
    the underlying <pic:spPr> element — the standard OOXML recipe for this,
    same category of manual XML patch as _add_hyperlink/_set_cell_background
    above."""
    pic = inline_shape._inline.graphic.graphicData.pic
    spPr = pic.spPr
    for tag in ("a:prstGeom", "a:custGeom"):
        existing = spPr.find(qn(tag))
        if existing is not None:
            spPr.remove(existing)
    geom = OxmlElement("a:prstGeom")
    geom.set("prst", "ellipse")
    geom.append(OxmlElement("a:avLst"))
    spPr.append(geom)


def _add_circular_photo(container, photo_data_uri: str, size: float):
    """Adds a square, circularly-clipped photo to `container` (a Document or
    a table _Cell — both support add_paragraph()), centered in its own
    paragraph. No-op if the photo can't be decoded."""
    stream = _decode_photo(photo_data_uri)
    if not stream:
        return None
    paragraph = container.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    picture = run.add_picture(stream, width=Inches(size), height=Inches(size))
    _make_picture_circular(picture)
    return paragraph


def _add_heading(container, text, color_hex, size=11, bold=True):
    heading = container.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(text.upper())
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color_hex)
    _add_bottom_border(heading, color_hex)
    return heading


def _add_label_value_rows(container, rows):
    table = container.add_table(rows=0, cols=2)
    table.autofit = True
    for label, value in rows:
        row = table.add_row()
        label_cell, value_cell = row.cells
        label_run = label_cell.paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(10)
        value_run = value_cell.paragraphs[0].add_run(value)
        value_run.font.size = Pt(10)


def _add_bullet(container, text):
    bullet = container.add_paragraph()
    bullet.paragraph_format.left_indent = Pt(18)
    bullet.paragraph_format.first_line_indent = Pt(-14)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.add_run(f"•\t{text}")


def _write_header_text(container, vm, theme_color, align_center=True):
    if vm["full_name"]:
        name_p = container.add_paragraph()
        if align_center:
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(vm["full_name"])
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.color.rgb = _rgb(theme_color)

    if vm["address"]:
        address_p = container.add_paragraph()
        if align_center:
            address_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_p.paragraph_format.space_after = Pt(0)
        address_run = address_p.add_run(vm["address"])
        address_run.font.size = Pt(9.5)
        address_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if vm["phone"] or vm["email"]:
        contact_p = container.add_paragraph()
        if align_center:
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(0)
        if vm["phone"]:
            contact_p.add_run(vm["phone"]).font.size = Pt(9.5)
        if vm["phone"] and vm["email"]:
            contact_p.add_run("   |   ").font.size = Pt(9.5)
        if vm["email"]:
            _add_hyperlink(contact_p, f"mailto:{vm['email']}", vm["email"], theme_color)

    if vm["link_items"]:
        links_p = container.add_paragraph()
        if align_center:
            links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, link in enumerate(vm["link_items"]):
            if i > 0:
                links_p.add_run("   |   ").font.size = Pt(9.5)
            _add_hyperlink(links_p, link["href"], link["text"], theme_color)


def _add_education_table(container, education_rows, theme_color):
    headers = ["Award", "Institute", "Address", "Percentage/Grade", "Duration"]
    table = container.add_table(rows=1, cols=len(headers))
    table.autofit = True
    header_row = table.rows[0]
    for cell, header in zip(header_row.cells, headers):
        _set_cell_background(cell, theme_color)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = _rgb("#FFFFFF")
    for edu in education_rows:
        row = table.add_row()
        values = [edu["degree"], edu["institute"], edu["address"], edu["percentage_grade"], edu["duration"]]
        for cell, value in zip(row.cells, values):
            run = cell.paragraphs[0].add_run(value)
            run.font.size = Pt(9.5)


def _add_experience_entries(container, experience_rows):
    for exp in experience_rows:
        entry = container.add_paragraph()
        entry.paragraph_format.space_after = Pt(0)
        company_run = entry.add_run(exp["company"])
        company_run.bold = True

        if exp["position"]:
            position_p = container.add_paragraph()
            position_p.paragraph_format.space_after = Pt(0)
            position_p.add_run(exp["position"]).italic = True

        if exp["tenure"]:
            tenure_p = container.add_paragraph()
            tenure_p.paragraph_format.space_after = Pt(2)
            tenure_run = tenure_p.add_run(f"Tenure: {exp['tenure']}")
            tenure_run.font.size = Pt(9)
            tenure_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        for bullet in exp["responsibilities"]:
            _add_bullet(container, bullet)


def _add_reference_entries(container, reference_rows):
    for ref in reference_rows:
        entry = container.add_paragraph()
        entry.paragraph_format.space_after = Pt(0)
        entry.add_run(ref["name"]).bold = True

        role = " – ".join(part for part in (ref["position"], ref["company"]) if part)
        if role:
            role_p = container.add_paragraph()
            role_p.paragraph_format.space_after = Pt(0)
            role_p.add_run(role).italic = True

        contact = "   |   ".join(part for part in (ref["phone"], ref["email"]) if part)
        if contact:
            contact_p = container.add_paragraph()
            contact_p.paragraph_format.space_after = Pt(6)
            contact_run = contact_p.add_run(contact)
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_declaration_section(document, vm, theme_color):
    """Always appended straight to `document` (never to a template's sidebar/
    main cell) so it lands after every other section regardless of layout —
    see the four _docx_* builders below, each of which calls this as its
    final statement. Deliberately doesn't reuse _add_heading: that always
    draws a full-width bottom border, but here the "line" is tied
    specifically to the date (only drawn when declaration_date is set) and
    is short, not a full-width section divider.
    """
    if not (vm["include_declaration"] and vm["declaration_text"]):
        return

    declaration_date = vm["declaration_date"]

    if declaration_date:
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        heading_cell, date_cell = table.rows[0].cells
        table.columns[0].width = Inches(1.7)
        table.columns[1].width = Inches(5.6)
        heading_cell.width = Inches(1.7)
        date_cell.width = Inches(5.6)

        heading_p = heading_cell.paragraphs[0]
        heading_p.paragraph_format.space_before = Pt(12)
        heading_run = heading_p.add_run("DECLARATION:")
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = _rgb(theme_color)

        date_p = date_cell.paragraphs[0]
        date_p.paragraph_format.space_before = Pt(12)
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_bottom_border(date_p, theme_color, size=8)
        date_run = date_p.add_run(declaration_date)
        date_run.font.size = Pt(9.5)
        date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        heading_p = document.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(12)
        heading_run = heading_p.add_run("DECLARATION:")
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        heading_run.font.color.rgb = _rgb(theme_color)

    # Free text from a textarea — split on newlines the same way "skills" is,
    # so a multi-line declaration doesn't collapse into one run-on paragraph.
    lines = vm["declaration_text"].split("\n")
    first_p = document.add_paragraph(lines[0])
    first_p.paragraph_format.space_before = Pt(4)
    for line in lines[1:]:
        document.add_paragraph(line)


def _configure_document(document):
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)


def _docx_classic(document, vm, theme_color):
    photo = vm["photo"]
    if photo:
        header_table = document.add_table(rows=1, cols=2)
        header_table.autofit = False
        text_cell, photo_cell = header_table.rows[0].cells
        header_table.columns[0].width = Inches(5.7)
        header_table.columns[1].width = Inches(1.3)
        text_cell.width = Inches(5.7)
        photo_cell.width = Inches(1.3)
        _write_header_text(text_cell, vm, theme_color, align_center=True)
        _add_circular_photo(photo_cell, photo, size=1.1)
    else:
        _write_header_text(document, vm, theme_color, align_center=True)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(4)
    _add_bottom_border(rule, theme_color, size=16)

    if vm["summary"]:
        _add_heading(document, "Summary", theme_color)
        document.add_paragraph(vm["summary"])

    if vm["personal_rows"]:
        _add_heading(document, "Personal Details", theme_color)
        _add_label_value_rows(document, vm["personal_rows"])

    if vm["passport_rows"]:
        _add_heading(document, "Passport Details", theme_color)
        _add_label_value_rows(document, vm["passport_rows"])

    if vm["physical_details_rows"]:
        _add_heading(document, "Physical Details", theme_color)
        _add_label_value_rows(document, vm["physical_details_rows"])

    if vm["emergency_contact_rows"]:
        _add_heading(document, "Emergency Contact", theme_color)
        _add_label_value_rows(document, vm["emergency_contact_rows"])

    if vm["preferred_position"]:
        _add_heading(document, "Preferred Position", theme_color)
        document.add_paragraph(vm["preferred_position"])

    if vm["medical_fitness"]:
        _add_heading(document, "Medical Fitness Status", theme_color)
        document.add_paragraph(vm["medical_fitness"])

    if vm["education_rows"]:
        _add_heading(document, "Academic Qualification", theme_color)
        _add_education_table(document, vm["education_rows"], theme_color)

    if vm["experience_rows"]:
        _add_heading(document, "Work Experience", theme_color)
        _add_experience_entries(document, vm["experience_rows"])

    if vm["language_lines"]:
        _add_heading(document, "Languages", theme_color)
        for line in vm["language_lines"]:
            _add_bullet(document, line)

    if vm["skills"]:
        _add_heading(document, "Skills", theme_color)
        for line in vm["skills"].split("\n"):
            document.add_paragraph(line)

    if vm["reference_rows"]:
        _add_heading(document, "References", theme_color)
        _add_reference_entries(document, vm["reference_rows"])

    _add_declaration_section(document, vm, theme_color)


def _docx_modern(document, vm, theme_color):
    white = "#FFFFFF"
    sidebar_width = Inches(2.2)
    main_width = Inches(4.8)

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    sidebar_cell, main_cell = table.rows[0].cells
    table.columns[0].width = sidebar_width
    table.columns[1].width = main_width
    sidebar_cell.width = sidebar_width
    main_cell.width = main_width
    _set_cell_background(sidebar_cell, theme_color)

    def sidebar_line(text, size=9.5, bold=False):
        p = sidebar_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = _rgb(white)

    def sidebar_bullet(text, size=9.5):
        p = sidebar_cell.add_paragraph()
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.first_line_indent = Pt(-10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"•\t{text}")
        run.font.size = Pt(size)
        run.font.color.rgb = _rgb(white)

    if vm["photo"]:
        photo_p = _add_circular_photo(sidebar_cell, vm["photo"], size=1.3)
        if photo_p is not None:
            photo_p.paragraph_format.space_after = Pt(8)

    if vm["full_name"]:
        name_p = sidebar_cell.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_p.paragraph_format.space_after = Pt(8)
        name_run = name_p.add_run(vm["full_name"])
        name_run.bold = True
        name_run.font.size = Pt(14)
        name_run.font.color.rgb = _rgb(white)

    contact_lines = [line for line in (vm["phone"], vm["email"], vm["address"]) if line]
    if contact_lines or vm["link_items"]:
        _add_heading(sidebar_cell, "Contact", white, size=10)
        for line in contact_lines:
            sidebar_line(line)
        for link in vm["link_items"]:
            sidebar_line(link["text"])

    if vm["physical_details_rows"]:
        _add_heading(sidebar_cell, "Physical Details", white, size=10)
        for label, value in vm["physical_details_rows"]:
            sidebar_line(f"{label}: {value}")

    if vm["emergency_contact_rows"]:
        _add_heading(sidebar_cell, "Emergency Contact", white, size=10)
        for label, value in vm["emergency_contact_rows"]:
            sidebar_line(f"{label}: {value}")

    if vm["preferred_position"]:
        _add_heading(sidebar_cell, "Preferred Position", white, size=10)
        sidebar_line(vm["preferred_position"])

    if vm["medical_fitness"]:
        _add_heading(sidebar_cell, "Medical Fitness Status", white, size=10)
        sidebar_line(vm["medical_fitness"])

    if vm["education_rows"]:
        _add_heading(sidebar_cell, "Education", white, size=10)
        for edu in vm["education_rows"]:
            if edu["degree"]:
                sidebar_line(edu["degree"], bold=True)
            if edu["institute"]:
                sidebar_line(edu["institute"], size=9)
            if edu["duration"]:
                sidebar_line(edu["duration"], size=8.5)

    if vm["skills"]:
        _add_heading(sidebar_cell, "Skills", white, size=10)
        for line in vm["skills"].split("\n"):
            if line.strip():
                sidebar_line(line.strip())

    if vm["language_lines"]:
        _add_heading(sidebar_cell, "Languages", white, size=10)
        for line in vm["language_lines"]:
            sidebar_bullet(line)

    if vm["summary"]:
        _add_heading(main_cell, "Summary", theme_color)
        main_cell.add_paragraph(vm["summary"])

    if vm["experience_rows"]:
        _add_heading(main_cell, "Experience", theme_color)
        _add_experience_entries(main_cell, vm["experience_rows"])

    # Added to `document`, not `main_cell` — the sidebar/main table is a
    # fixed decorative layout, so this deliberately lands after the whole
    # table rather than inside one of its cells, to stay the true last thing
    # in the document.
    _add_declaration_section(document, vm, theme_color)


def _add_light_heading(container, text, color_hex, size=10.5):
    heading = container.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(text.upper())
    run.font.size = Pt(size)
    run.font.bold = False
    run.font.color.rgb = _rgb(color_hex)
    return heading


def _docx_minimalist(document, vm, theme_color):
    if vm["photo"]:
        header_table = document.add_table(rows=1, cols=2)
        header_table.autofit = False
        photo_cell, text_cell = header_table.rows[0].cells
        header_table.columns[0].width = Inches(1.0)
        header_table.columns[1].width = Inches(6.0)
        photo_cell.width = Inches(1.0)
        text_cell.width = Inches(6.0)
        _add_circular_photo(photo_cell, vm["photo"], size=0.8)
        _write_header_text(text_cell, vm, theme_color, align_center=False)
    else:
        _write_header_text(document, vm, theme_color, align_center=False)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after = Pt(8)
    _add_bottom_border(rule, "#DDDDDD", size=6)

    if vm["summary"]:
        _add_light_heading(document, "Summary", theme_color)
        p = document.add_paragraph(vm["summary"])
        p.paragraph_format.space_after = Pt(8)

    if vm["physical_details_rows"]:
        _add_light_heading(document, "Physical Details", theme_color)
        _add_label_value_rows(document, vm["physical_details_rows"])

    if vm["emergency_contact_rows"]:
        _add_light_heading(document, "Emergency Contact", theme_color)
        _add_label_value_rows(document, vm["emergency_contact_rows"])

    if vm["preferred_position"]:
        _add_light_heading(document, "Preferred Position", theme_color)
        document.add_paragraph(vm["preferred_position"])

    if vm["medical_fitness"]:
        _add_light_heading(document, "Medical Fitness Status", theme_color)
        document.add_paragraph(vm["medical_fitness"])

    if vm["experience_rows"]:
        _add_light_heading(document, "Experience", theme_color)
        _add_experience_entries(document, vm["experience_rows"])

    if vm["education_rows"]:
        _add_light_heading(document, "Education", theme_color)
        for edu in vm["education_rows"]:
            if edu["degree"]:
                line = document.add_paragraph()
                line.paragraph_format.space_after = Pt(0)
                line.add_run(edu["degree"]).bold = True
            if edu["institute"]:
                sub = document.add_paragraph()
                sub.paragraph_format.space_after = Pt(0)
                sub.add_run(edu["institute"]).italic = True
            if edu["duration"]:
                dur = document.add_paragraph()
                dur.paragraph_format.space_after = Pt(8)
                dur_run = dur.add_run(edu["duration"])
                dur_run.font.size = Pt(9)
                dur_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if vm["language_lines"]:
        _add_light_heading(document, "Languages", theme_color)
        for line in vm["language_lines"]:
            _add_bullet(document, line)

    if vm["skills"]:
        _add_light_heading(document, "Skills", theme_color)
        for line in vm["skills"].split("\n"):
            document.add_paragraph(line)

    _add_declaration_section(document, vm, theme_color)


def _docx_academic(document, vm, theme_color):
    normal = document.styles["Normal"]
    normal.font.name = "Georgia"

    _write_header_text(document, vm, theme_color, align_center=True)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(6)
    _add_bottom_border(rule, theme_color, size=12)

    if vm["physical_details_rows"]:
        _add_heading(document, "Physical Details", theme_color)
        _add_label_value_rows(document, vm["physical_details_rows"])

    if vm["emergency_contact_rows"]:
        _add_heading(document, "Emergency Contact", theme_color)
        _add_label_value_rows(document, vm["emergency_contact_rows"])

    if vm["preferred_position"]:
        _add_heading(document, "Preferred Position", theme_color)
        document.add_paragraph(vm["preferred_position"])

    if vm["medical_fitness"]:
        _add_heading(document, "Medical Fitness Status", theme_color)
        document.add_paragraph(vm["medical_fitness"])

    if vm["education_rows"]:
        _add_heading(document, "Education", theme_color)
        _add_education_table(document, vm["education_rows"], theme_color)

    if vm["publication_rows"]:
        _add_heading(document, "Publications", theme_color)
        for pub in vm["publication_rows"]:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            title_run = p.add_run(pub["title"])
            title_run.italic = True
            rest = ", ".join(part for part in (pub["venue"], pub["year"]) if part)
            if rest:
                p.add_run(f" — {rest}")

    if vm["experience_rows"]:
        _add_heading(document, "Research / Professional Experience", theme_color)
        _add_experience_entries(document, vm["experience_rows"])

    if vm["language_lines"]:
        _add_heading(document, "Languages", theme_color)
        for line in vm["language_lines"]:
            _add_bullet(document, line)

    if vm["skills"]:
        _add_heading(document, "Skills", theme_color)
        for line in vm["skills"].split("\n"):
            document.add_paragraph(line)

    _add_declaration_section(document, vm, theme_color)


DOCX_BUILDERS = {
    "classic": _docx_classic,
    "modern": _docx_modern,
    "minimalist": _docx_minimalist,
    "academic": _docx_academic,
}


def render_cv_docx(cv) -> bytes:
    vm = build_cv_view_model(cv)
    theme_color = vm["theme_color"]

    document = Document()
    _configure_document(document)

    builder = DOCX_BUILDERS.get(vm["template"], _docx_classic)
    builder(document, vm, theme_color)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
